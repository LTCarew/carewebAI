"""
Assemble the captured CareWeb AI journeys into a plain-language Word guide.

Run after generate_screenshots.py:
    venv/Scripts/python.exe generate_journey_walkthrough.py

Input:
    screenshots/manifest.json and the referenced PNG files

Output:
    CareWebAI_User_Journey_Walkthrough.docx
"""

import json
import os
from datetime import datetime
from pathlib import Path


BASE_DIR = Path(__file__).resolve().parent
SCREENSHOT_DIR = BASE_DIR / "screenshots"
OUTPUT_PATH = BASE_DIR / "CareWebAI_User_Journey_Walkthrough.docx"


def build_document():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    manifest_path = SCREENSHOT_DIR / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(
            "screenshots/manifest.json not found. Run generate_screenshots.py first."
        )
    with manifest_path.open(encoding="utf-8") as handle:
        entries = json.load(handle)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    def set_cell_shading(cell, fill="E8EDF8"):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt({1: 16, 2: 13, 3: 11}.get(level, 11))
        return p

    def paragraph(text="", bold=False, italic=False, size=10.5):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run.font.size = Pt(size)
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    def table(headers, rows, font_size=9):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for index, value in enumerate(headers):
            cell = t.rows[0].cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            set_cell_shading(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for row in rows:
            cells = t.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = ""
                run = cells[index].paragraphs[0].add_run(str(value))
                run.font.name = "Arial"
                run.font.size = Pt(font_size)
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        doc.add_paragraph()
        return t

    # ------------------------------------------------------------------ cover
    doc.add_paragraph()
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("CareWeb AI")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)

    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_subtitle.add_run("User Journey Walkthrough")
    run.font.name = "Arial"
    run.font.size = Pt(18)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(
        f"Careworker, Staff Admin, and Client Workflows — {datetime.now():%B %d, %Y}"
    )
    run.font.name = "Arial"
    run.font.size = Pt(12)

    doc.add_paragraph()
    cover_meta = doc.add_paragraph()
    cover_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_meta.add_run(
        "Prepared by: CareWeb AI Development Team\n"
        "Organization: Center for Independent Living (CIL)\n"
        "Demonstration environment: CIL-Care seeded demo data\n"
        "Capture format: 1440px-wide full-page headless Chrome screenshots"
    )
    run.font.name = "Arial"
    run.font.size = Pt(11)
    doc.add_page_break()

    # --------------------------------------------------------------- orientation
    heading("How CareWeb AI Fits Together")
    paragraph(
        "CareWeb AI is a Personal Attendant Services coordination workspace. It brings "
        "together people receiving support, careworkers, and staff at an organization such "
        "as an Independent Living Center. Each person sees a role-specific workspace, while "
        "the underlying relationship can move through matching, approvals, scheduling, "
        "sessions, feedback, and human follow-up."
    )
    paragraph(
        "The screenshots in this guide use a demonstration organization called CIL-Care. "
        "The records are seeded demo records, not real participant data. Names, contact "
        "details, ratings, and schedules are included only to make the workflow look and "
        "behave like a populated application."
    )
    table(
        ["Role", "Primary responsibility", "Main features shown"],
        [
            ("Careworker", "Describe skills and availability, review relationships and schedules, and share experience feedback.", "Dashboard, profile, schedule, rating workflow"),
            ("Staff admin", "Coordinate the organization, review participants, evaluate matching suggestions, and provide relationship support.", "Organization dashboard, pools, AI-assisted matching, Stability Snapshot"),
            ("Client / person receiving support", "Describe care needs and preferences, review matches and schedules, and share feedback about the relationship.", "Dashboard, profile, schedule, rating workflow"),
        ],
        font_size=9,
    )
    heading("End-to-end workflow", level=2)
    for item in [
        "A client and careworker provide structured information such as care needs, experience, availability, location, language, transportation, and preferences.",
        "The matching workflow compares those attributes and provides a compatibility score with a factor breakdown and human-readable reasoning.",
        "Staff review potential pairings and can approve or coordinate a match; the AI output supports review and does not make care decisions by itself.",
        "The parties use a shared schedule with recurring time slots and explicit approval states.",
        "After sessions, the client and careworker independently rate care fit, communication, reliability, and workload balance.",
        "The Stability Snapshot summarizes patterns such as baseline ratings, trends, disagreement, single-metric concerns, and travel burden so staff can decide whether a supportive check-in is appropriate.",
        "A staff member can flag a relationship for stabilization review. The flag creates a human follow-up task; it is not an automatic termination or care decision.",
    ]:
        bullet(item)
    paragraph(
        "Important interpretation note: AI-assisted matching and stability indicators are "
        "decision-support features. Human review remains part of the workflow, and the "
        "participants' circumstances and preferences remain central to coordination.",
        italic=True,
    )
    doc.add_page_break()

    # ------------------------------------------------------------- persona pages
    persona_info = {
        "careworker": (
            "Careworker Journey",
            "This journey shows how a careworker enters the system, reviews the information used for matching, checks a live schedule, and contributes independent feedback about the working relationship.",
        ),
        "staff_admin": (
            "Staff Admin Journey",
            "This journey shows how organization staff move from a consolidated dashboard to participant review, explainable matching, relationship monitoring, and supportive escalation.",
        ),
        "client": (
            "Client Journey",
            "This journey shows how a person receiving support describes their needs, reviews the coordination workspace, checks the agreed schedule, and shares feedback in their own voice.",
        ),
    }

    for persona in ("careworker", "staff_admin", "client"):
        title, overview = persona_info[persona]
        heading(title)
        paragraph(overview)
        persona_entries = sorted(
            [entry for entry in entries if entry["persona"] == persona],
            key=lambda entry: entry["step"],
        )
        for entry in persona_entries:
            heading(f"Step {entry['step']}: {entry['title']}", level=2)
            paragraph(f"What this step does: {entry['purpose']}", bold=True)
            paragraph(entry["feature_summary"])
            image_path = BASE_DIR / entry["screenshot"]
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(image_path), width=Inches(6.65))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(
                    f"Figure {persona_entries.index(entry) + 1}. {entry['title']} — {entry['purpose']}"
                )
                run.italic = True
                run.font.name = "Arial"
                run.font.size = Pt(9)
            else:
                paragraph(f"Screenshot unavailable at {entry['screenshot']}", italic=True)
            paragraph(
                "Reviewer takeaway: this screen is part of the role-specific workflow and "
                "is designed to keep coordination information visible, understandable, and "
                "available for human review.",
                italic=True,
                size=9.5,
            )
            # Avoid putting every step on a new page; Word may naturally break large
            # screenshots while keeping the caption close to its image.
        doc.add_page_break()

    # --------------------------------------------------------------- comparison
    heading("Comparing the Three Journeys")
    paragraph(
        "The application uses the same underlying relationship data while presenting it "
        "differently according to the user's role. This keeps participants focused on their "
        "own responsibilities and gives staff the broader coordination view needed for support."
    )
    table(
        ["Capability", "Careworker view", "Staff admin view", "Client view"],
        [
            ("Profile information", "Skills, experience, availability, language, transportation, and preferences.", "Reviews participant information for coordination and matching.", "Care needs, programs, availability, language, preferences, and context."),
            ("Matching", "Can review suggested relationships relevant to the careworker.", "Can compare candidates, inspect factor breakdowns, and use AI-assisted suggestions.", "Can review careworker suggestions relevant to the client's needs."),
            ("Scheduling", "Reviews assigned schedules and responds to session details.", "Coordinates schedules and sees approval/relationship status across the organization.", "Creates or reviews schedules and confirms the agreed support arrangement."),
            ("Feedback", "Rates care fit, communication, reliability, and workload.", "Reviews aggregated patterns and decides whether a check-in is useful.", "Rates the same dimensions independently from the client's perspective."),
            ("Stability support", "Provides information that may reveal a need for support.", "Sees explainable signals and can flag a relationship for human follow-up.", "Feedback contributes to a fuller picture without an automated decision about care."),
        ],
        font_size=8,
    )
    heading("What the AI features mean", level=2)
    for item in [
        "The criteria-based tag view is a transparent first step: it shows who shares a selected skill or care need without ranking people.",
        "AI-assisted matching uses structured profile information to rank compatibility and explain strengths or concerns; it is not a replacement for participant choice or staff judgment.",
        "The local scoring layer preserves a deterministic factor breakdown such as care-tag overlap, availability, location, transportation, and language. A ChatGPT-enhanced narrative can make that result easier to interpret while the stored breakdown preserves traceability.",
        "The AI capture identifies the narrative source in the interface. If the external service is unavailable, the product falls back to the local scoring model rather than failing silently.",
        "Stability indicators summarize participant feedback over time. They are signals for conversation and support, not labels assigned to a person and not an automatic action.",
    ]:
        bullet(item)

    heading("Demonstration and privacy notes", level=2)
    paragraph(
        "This walkthrough was generated from an isolated test database populated with synthetic "
        "demo records. It is intended to demonstrate navigation, features, and role separation. "
        "When presenting the application publicly, continue to use synthetic or de-identified "
        "records and do not expose real participant contact information in screenshots."
    )
    paragraph(
        "The capture process used a 1440px-wide headless Chrome browser and saved full-page original "
        "PNG files beside this document. The accompanying manifest records each screenshot's "
        "route, purpose, and plain-language feature summary so the walkthrough can be regenerated "
        "after future interface changes.",
        italic=True,
        size=9.5,
    )

    doc.core_properties.title = "CareWeb AI User Journey Walkthrough"
    doc.core_properties.subject = "Careworker, staff admin, and client application workflows"
    doc.core_properties.author = "CareWeb AI Development Team"
    doc.core_properties.comments = "Generated from synthetic CIL-Care demo data."
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved: {path}")