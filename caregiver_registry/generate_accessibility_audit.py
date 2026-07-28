
"""
CareWeb — Accessibility Audit Document Generator
====================================================

Creates a grant-ready Word document documenting the CareWeb accessibility
review.  The document follows the same general format as
``CareWebAI_TRL3_Submission.docx``: cover page, numbered sections, tables,
bulleted findings, and appendices.

This generator intentionally reports the audit performed rather than claiming
that an automated axe scan or assistive-technology usability study was run.
The source review is a WCAG 2.1 AA-oriented static review of the templates,
CSS, and JavaScript in the repository.  The document identifies the additional
testing needed before making a formal conformance claim.

Usage (from caregiver_registry/):
    venv/Scripts/python.exe generate_accessibility_audit.py

Output:
    CareWeb_Accessibility_Audit.docx
"""

from datetime import datetime
import os


def build_document(output_path):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Match the margins and restrained typography used by the TRL-3 generator.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    now_str = datetime.now().strftime("%B %d, %Y")

    def set_cell_shading(cell, fill="E8EDF8"):
        properties = cell._tc.get_or_add_tcPr()
        shading = properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            properties.append(shading)
        shading.set(qn("w:fill"), fill)

    def set_cell_text(cell, text, bold=False, size=10):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(text))
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    def table(headers, rows, widths=None, font_size=9):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for index, header in enumerate(headers):
            set_cell_text(t.rows[0].cells[index], header, bold=True, size=font_size)
            set_cell_shading(t.rows[0].cells[index])
        for row in rows:
            cells = t.add_row().cells
            for index, value in enumerate(row):
                set_cell_text(cells[index], value, size=font_size)
        if widths:
            for row in t.rows:
                for index, width in enumerate(widths):
                    row.cells[index].width = Inches(width)
        doc.add_paragraph()
        return t

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        run = h.runs[0] if h.runs else h.add_run()
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
        return h

    def body(text="", bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run.font.size = Pt(size)
        return p

    def bullet(text, level=0):
        style = "List Bullet" if level == 0 else "List Bullet 2"
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        return p

    def numbered(text):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        return p

    def mono(text, size=9):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        return p

    def page_break():
        doc.add_page_break()

    # ---------------------------------------------------------------------
    # COVER
    # ---------------------------------------------------------------------
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CareWeb")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Accessibility Audit")
    run.font.name = "Arial"
    run.font.size = Pt(16)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run(f"WCAG 2.1 AA Code Review — {now_str}")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Prepared by: CareWeb Development Team\n"
        "Organization: Center for Independent Living (CIL)\n"
        "Product: Personal Care Coordination and Stabilization\n"
        "Review type: Static source-code accessibility audit"
    )
    run.font.name = "Arial"
    run.font.size = Pt(11)
    page_break()

    # ---------------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # ---------------------------------------------------------------------
    heading("1. Executive Summary")
    body(
        "CareWeb was reviewed against the principles and applicable success "
        "criteria of the Web Content Accessibility Guidelines (WCAG) 2.1 Level AA. "
        "The review covered the shared layout, navigation, authentication, application "
        "forms, dashboards, matching tables, stability monitoring, rating workflow, "
        "CSS, and JavaScript interactions present in the repository."
    )
    body(
        "Overall assessment: PARTIAL CONFORMANCE / REMEDIATION REQUIRED. The CareWeb application "
        "already contains several accessibility-supportive design decisions, including "
        "semantic landmarks, explicit form labels, keyboard-visible focus styling, "
        "table header scopes in major dashboard views, text labels alongside status colors, "
        "accessible font sizing, and several useful ARIA attributes. The review also found "
        "important gaps that should be addressed before making a formal WCAG 2.1 AA "
        "conformance statement."
    )

    table(
        ["Assessment area", "Current assessment", "Primary evidence"],
        [
            ("Semantic structure", "Mostly strong", "base.html landmarks; headings and table structures"),
            ("Forms and errors", "Partial", "Explicit labels present; error/help associations need strengthening"),
            ("Keyboard operation", "Partial", "Focus styling and burger button present; dropdown and popover behavior need testing/fixes"),
            ("Color and contrast", "Mixed", "Most core text passes; some documented tokens fail or are borderline for WCAG AA"),
            ("Dynamic content", "Partial", "Live region exists; step changes and loading overlay need announcements/focus handling"),
            ("Tables and data views", "Mostly strong", "Major tables use scope; rating history table needs explicit scopes/names"),
            ("Motion and preferences", "Needs remediation", "Spinner animation has no reduced-motion alternative"),
            ("Assistive-technology validation", "Not yet performed", "Requires screen-reader, keyboard-only, zoom, and switch-access testing"),
        ],
        widths=[1.45, 1.35, 3.8],
    )

    heading("1.1 Key strengths", level=2)
    for item in [
        "The document declares <html lang=\"en\"> and a responsive viewport in base.html.",
        "The shared layout provides navigation, a main content area, and a footer; the navbar has an accessible navigation label.",
        "The application uses explicit label-for associations for many manually rendered form fields and provides visible focus styles.",
        "The AI loading overlay uses role=\"status\" and aria-live=\"polite\", and the navbar button uses aria-expanded and aria-controls.",
        "Major dashboard and stability tables use scope=\"col\" / scope=\"row\" and accessible region/table labels.",
        "Stability indicators do not rely on color alone: each status includes a visible text label and a decorative dot.",
        "The CSS uses a relatively large default font size and Atkinson Hyperlegible, which supports readability and low-vision use.",
    ]:
        bullet(item)

    heading("1.2 Highest-priority actions", level=2)
    for item in [
        "Add a skip link and a programmatic target on the main content area.",
        "Make multi-step form state programmatically available: identify the current step, announce changes, preserve focus, and provide clear validation/error associations.",
        "Repair the organization dropdown and custom information-tip semantics for reliable keyboard and screen-reader use.",
        "Add reduced-motion behavior for the AI spinner and smooth scrolling.",
        "Correct or confirm contrast for the documented failing color combinations, especially small muted text, danger text, and warning backgrounds.",
        "Complete manual testing with people who use screen readers, keyboard-only navigation, magnification, voice input, and switch access.",
    ]:
        bullet(item)

    page_break()

    # ---------------------------------------------------------------------
    # 2. SCOPE AND METHODOLOGY
    # ---------------------------------------------------------------------
    heading("2. Scope and Methodology")
    body(
        "This audit is a source-code review intended to identify accessibility risks and "
        "document a practical remediation plan for grant and product-development purposes. "
        "It is not a legal determination, a VPAT/ACR, or a final certification of WCAG "
        "conformance. Findings should be re-tested after remediation."
    )
    heading("2.1 Standard and review approach", level=2)
    for item in [
        "Reference standard: WCAG 2.1 Level AA, including perceivable, operable, understandable, and robust principles.",
        "Review method: manual inspection of rendered HTML templates, CSS tokens/rules, and JavaScript behavior.",
        "Evidence method: findings cite repository paths and line ranges as observed during this review.",
        "Contrast method: representative foreground/background token pairs were evaluated using WCAG relative-luminance calculations; gradients, hover states, and all browser rendering contexts still require live verification.",
        "Status terms: PASS means the reviewed implementation shows a strong control; PARTIAL means a control exists but has a material gap; NEEDS REMEDIATION means a likely WCAG issue or unverified critical behavior was identified.",
    ]:
        bullet(item)

    heading("2.2 Scope reviewed", level=2)
    table(
        ["Area", "Files / components reviewed"],
        [
            ("Shared application shell", "templates/base.html; templates/navbar.html"),
            ("Public and authentication pages", "templates/home.html; templates/registration/login.html"),
            ("Application forms", "templates/registry/caregiver_apply.html; client_apply.html; schedule_form.html; schedule_entry_rate.html"),
            ("Dashboards and data tables", "templates/registry/org_dashboard.html; caregiver_dashboard.html; client_dashboard.html; matching/_match_table.html"),
            ("AI and stability views", "templates/matching/stability_detail.html; matching/_score_cell.html; matching/ai_match_*.html"),
            ("Visual and interaction layer", "static/css/style.css; inline JavaScript in base.html, caregiver_apply.html, stability_detail.html, schedule_entry_rate.html"),
        ],
        widths=[1.8, 4.8],
    )

    heading("2.3 Testing limitations", level=2)
    for item in [
        "No live automated axe-core, Lighthouse, or pa11y result is represented in this document. A browser-based automated scan should be added to the next test cycle.",
        "No user testing with a screen reader, magnification software, voice control, or switch device was performed as part of this code review.",
        "A static review cannot confirm actual tab order, focus visibility at all viewport sizes, browser-specific native control behavior, or the accessibility tree produced by a particular browser/assistive-technology combination.",
        "Color contrast results are representative token checks, not a substitute for checking every rendered state, gradient, image, focus state, or user-selected high-contrast mode.",
    ]:
        bullet(item)

    page_break()

    # ---------------------------------------------------------------------
    # 3. FINDINGS
    # ---------------------------------------------------------------------
    heading("3. Findings by Accessibility Category")
    body(
        "The findings below separate controls that are already present from controls that "
        "need strengthening. File references are intended to make remediation and later "
        "verification traceable."
    )

    findings = [
        (
            "3.1 Semantic structure and landmarks",
            "PARTIAL",
            "WCAG 1.3.1; 2.4.1; 2.4.6",
            [
                "PASS: base.html provides a navigation include, a main element, and a footer; navbar.html labels the navigation with aria-label=\"Main Navigation\".",
                "GAP: there is no skip-to-main-content link, and the main element has no id/tabindex target for keyboard users who need to bypass repeated navigation (base.html:27–55).",
                "GAP: the hero is a section-like banner but is not identified as a complementary/banner landmark; this is not necessarily a failure, but landmark clarity could improve navigation for screen-reader users.",
            ],
            "Add a visually hidden skip link as the first focusable element and give main an id such as main-content with tabindex=\"-1\". Verify focus styling and return focus after activation.",
        ),
        (
            "3.2 Forms, labels, instructions, and errors",
            "PARTIAL",
            "WCAG 1.3.1; 1.3.5; 3.3.1; 3.3.2; 3.3.3; 4.1.2",
            [
                "PASS: caregiver_apply.html and login.html manually render explicit labels using each field's id_for_label (caregiver_apply.html:49–137; login.html:27–49).",
                "GAP: help text and field errors are visually adjacent but are not consistently connected with aria-describedby; errors are not consistently assigned an id that can be referenced by the input.",
                "GAP: form-level non-field errors are presented in a notification but are not consistently given role=\"alert\" or a focus target (caregiver_apply.html:33–37; login.html:21–25).",
                "GAP: novalidate is used on forms, so robust server-side validation feedback and accessible error recovery are essential.",
                "GAP: required fields show a visual asterisk with aria-label=\"required\", but the underlying widget's required/aria-required state should be verified for every field.",
            ],
            "Generate stable ids for help/error text, add aria-describedby and aria-invalid when errors exist, use a focused error summary after failed submission, and ensure server-side validation returns the user to the first invalid step.",
        ),
        (
            "3.3 Multi-step application wizard",
            "NEEDS REMEDIATION",
            "WCAG 1.3.1; 2.4.3; 2.4.6; 3.3.1; 4.1.2",
            [
                "The six progress steps are div elements with data-step values, but the active step is represented only by a CSS class; there is no aria-current or programmatic progress relationship (caregiver_apply.html:20–28; 347–360).",
                "The JavaScript changes sections with display:none and scrolls smoothly to the top, but does not move focus to the new section heading or announce the new step (caregiver_apply.html:347–363).",
                "Next advances without validating the current step in the browser, which can let a user move past required fields without immediate feedback (caregiver_apply.html:365–370).",
                "The hidden Previous/Next/Submit controls are visually toggled; the implementation should be tested to ensure hidden controls are not focusable in all target browsers.",
            ],
            "Use a fieldset/legend or equivalent step structure, mark the current progress item with aria-current=\"step\", announce \"Step X of 6\", move focus to the new step heading, validate before advancing, and provide an accessible summary of missing/invalid fields.",
        ),
        (
            "3.4 Navigation, dropdowns, and keyboard operation",
            "PARTIAL",
            "WCAG 2.1.1; 2.4.3; 2.4.7; 4.1.2",
            [
                "PASS: the mobile navbar button is a real button and exposes aria-expanded and aria-controls (navbar.html:7–18); base.html updates aria-expanded when toggled (base.html:65–79).",
                "GAP: the organization switcher uses a navbar-link without an href or explicit button semantics and relies on Bulma hover behavior (navbar.html:39–57). Keyboard opening, expanded state, and escape/return behavior require remediation or explicit verification.",
                "GAP: the mobile menu does not close on Escape, outside click, or navigation, and focus is not returned to the burger button after closing (base.html:65–79).",
                "GAP: generic interactive links and controls should be checked at 200% zoom and on narrow screens for clipping and reachable focus.",
            ],
            "Implement the organization switcher as a button with aria-haspopup=\"true\", aria-expanded, and a controlled menu; support Enter/Space/Escape and focus return. Add Escape handling and focus management to the mobile menu.",
        ),
        (
            "3.5 Color contrast and non-color communication",
            "MIXED",
            "WCAG 1.4.3; 1.4.11",
            [
                "Representative passing pairs: --app-primary #7e6b8f on white is approximately 4.79:1; --app-primary-dark #665575 is approximately 6.73:1; link blue #3c6f99 is approximately 5.35:1; body text #0f172a is approximately 17.85:1.",
                "Potential failures: muted section-label color #94a3b8 on white is approximately 2.56:1 and is not suitable for normal-sized text; bright warning token #f2e94e has approximately 1.27:1 against white and must not be used as a text/background combination without a dark foreground.",
                "The danger token #da3e52 is approximately 4.38:1 on white, slightly below the 4.5:1 normal-text threshold; use a darker danger text token or reserve it for large text/UI indicators.",
                "PASS: stability statuses include text labels in addition to colored dots, and the CSS comments explicitly state that status is not conveyed by color alone (style.css:749–752).",
            ],
            "Replace low-contrast muted text with a darker token, verify danger text in every rendered state, and test focus/hover/border contrast as non-text UI components. Keep a text label with every status color.",
        ),
        (
            "3.6 Focus visibility and target size",
            "PARTIAL",
            "WCAG 2.4.7; 2.4.11 (WCAG 2.2 reference); 2.5.5",
            [
                "PASS: global focus rules use a 3px outline and app-focus color for inputs, buttons, links, and navbar controls (style.css:258–267; 54–59).",
                "GAP: information-tip buttons are sized at 1.2em, approximately 21–22px at the default font size, which is below the 44×44 CSS pixel target recommended by WCAG 2.2's target-size criterion (style.css:1065–1081).",
                "GAP: focus management has not been verified when custom popovers open, when the mobile menu changes, or when the wizard changes steps.",
            ],
            "Increase interactive hit areas to at least 44×44px where practical, preserve a visible focus indicator on every custom control, and run keyboard-only tests through every workflow.",
        ),
        (
            "3.7 ARIA, tooltips, and dynamic announcements",
            "PARTIAL",
            "WCAG 4.1.2; 4.1.3; 1.3.1",
            [
                "PASS: the AI loading overlay has role=\"status\" and aria-live=\"polite\" (base.html:19–24).",
                "PASS: info-tip buttons expose an accessible label and update aria-expanded when toggled (stability_detail.html:117–123; 445–490).",
                "GAP: the tooltip content has role=\"tooltip\" but is not given a unique id and is not connected to the button with aria-describedby or aria-controls; screen readers may not reliably receive the explanation (stability_detail.html:117–123; style.css:1100–1139).",
                "GAP: changing a rating slider updates a visible value span but does not announce the updated value with an accessible relationship or live region (schedule_entry_rate.html:239–258).",
                "GAP: the loading overlay is displayed when an AI form is submitted but does not move focus into the status dialog, disable the initiating form, or provide a completion/error announcement.",
            ],
            "Give each tooltip a unique id and connect it with aria-describedby, or use a simpler disclosure pattern. Associate slider output with aria-valuetext/aria-describedby and announce wizard/loading completion states without creating excessive speech.",
        ),
        (
            "3.8 Tables and data relationships",
            "MOSTLY STRONG",
            "WCAG 1.3.1; 1.3.2; 2.4.6",
            [
                "PASS: org_dashboard.html tables use scope=\"col\" and wrap schedule data in a labeled table region (org_dashboard.html:96–107).",
                "PASS: matching/_match_table.html uses scope=\"col\" and a table/region aria-label (matching/_match_table.html:26–52).",
                "PASS: stability_detail.html uses scope=\"col\" and scope=\"row\" for the signal table and provides aria-labels for rating tables (stability_detail.html:105–112; 270–271).",
                "GAP: the past-ratings table in schedule_entry_rate.html does not declare scope attributes or an accessible name (schedule_entry_rate.html:200–212).",
                "GAP: status badges and emoji flags should be tested with a screen reader to ensure decorative symbols do not obscure the actual status text.",
            ],
            "Add scope=\"col\" to every rating-history header and an aria-label or caption describing the table. Verify column relationships and reading order with a screen reader at narrow widths.",
        ),
        (
            "3.9 Motion, animation, and scrolling",
            "NEEDS REMEDIATION",
            "WCAG 2.3.3 (AAA reference); WCAG 2.2.2; user preference support",
            [
                "The AI spinner continuously animates with @keyframes ai-spin and has no @media (prefers-reduced-motion: reduce) override (style.css:709–720).",
                "The multi-step wizard uses smooth scrolling on every step transition (caregiver_apply.html:362), which can be disorienting for users who disable motion or use magnification.",
                "The application should respect prefers-reduced-motion even though the specific WCAG 2.1 AA criterion is not generally a blanket requirement for this type of short loading animation.",
            ],
            "Add a reduced-motion media query that disables the spinner animation and changes smooth scrolling to instant or no scrolling. Test for vestibular sensitivity and screen magnification.",
        ),
        (
            "3.10 Icons, emoji, and visual meaning",
            "PARTIAL",
            "WCAG 1.1.1; 1.3.3; 1.4.1",
            [
                "PASS: decorative home-page icon badges are marked aria-hidden=\"true\" (home.html:39; 71).",
                "PASS: the stability status dot is marked aria-hidden=\"true\" while the status label remains visible (stability_detail.html:84–88; matching/_match_table.html:113–117).",
                "GAP: many headings and buttons begin with emoji (for example, Organization Dashboard, Active Matches, Rate Your Experience, and Flag for Stabilization Review). Screen-reader pronunciation varies by platform; the text label must remain complete without the emoji.",
                "GAP: the flag badge includes an emoji and a screen-reader-only phrase, but the visual title attribute is not a reliable accessible description (matching/_match_table.html:119–123).",
            ],
            "Wrap decorative emoji in aria-hidden spans or remove them where unnecessary. Ensure all meaning is carried by adjacent text and never by emoji or color alone.",
        ),
        (
            "3.11 Responsive layout and readability",
            "PARTIAL",
            "WCAG 1.4.4; 1.4.10; 1.4.12",
            [
                "PASS: the viewport is configured, the base font is 18px, and forms/tables include responsive CSS rules (base.html:5–8; style.css:22–24; 206–214; 656–662).",
                "GAP: stability panels and tooltip bubbles use fixed positioning and viewport-based widths on small screens; live testing is needed to confirm they do not cover essential content or trap focus (style.css:998–1011; 1142–1160).",
                "GAP: dense data tables may require horizontal scrolling; the region must have a clear accessible name and the scroll affordance should be discoverable at 200% zoom.",
            ],
            "Test at 200% and 400% zoom, 320 CSS-pixel width, increased text spacing, and browser text-only zoom. Ensure popovers, tables, and forms remain usable without loss of information or functionality.",
        ),
        (
            "3.12 Plain language and user control",
            "MOSTLY STRONG",
            "WCAG 3.1.1; 3.1.3; 3.2.1",
            [
                "PASS: the stability detail page includes a plain-language glossary and explains each rating metric before presenting the rating history (stability_detail.html:30–69).",
                "PASS: AI recommendations and stability indicators are presented for human review rather than being represented in the reviewed templates as autonomous care decisions.",
                "GAP: the AI overlay and multi-step wizard should provide explicit status/error text when an operation fails, times out, or cannot proceed, rather than relying on a visual state change.",
            ],
            "Retain plain-language explanations, add consistent operation failure and recovery messages, and user-test the terminology with people who direct their own care and PAS staff.",
        ),
    ]

    for title_text, status, criteria, evidence, recommendation in findings:
        heading(title_text, level=2)
        status_p = doc.add_paragraph()
        r = status_p.add_run(f"Assessment: {status}    |    Relevant criteria: {criteria}")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        for item in evidence:
            bullet(item)
        body("Recommended control:", bold=True)
        body(recommendation)

    page_break()

    # ---------------------------------------------------------------------
    # 4. PRIORITIZED REMEDIATION PLAN
    # ---------------------------------------------------------------------
    heading("4. Prioritized Remediation Plan")
    body(
        "Priorities reflect user impact, risk of blocking a workflow, and the effort needed "
        "to establish a reliable accessibility baseline. Priority 1 items should be addressed "
        "before a grant or production statement that implies WCAG AA readiness."
    )
    table(
        ["Priority", "Remediation", "Acceptance evidence"],
        [
            ("1", "Add skip link and main-content focus target.", "Keyboard user reaches main content in one action; focus is visible."),
            ("1", "Make wizard state accessible: aria-current, focus movement, announcements, step validation, and error summary.", "Keyboard and screen-reader tester can complete, correct, and resubmit both application wizards."),
            ("1", "Fix form error associations and invalid-state communication.", "Every error is announced and associated with its input; focus reaches the first error."),
            ("1", "Repair organization dropdown and mobile-menu keyboard behavior.", "Menu opens/closes with keyboard; Escape works; focus does not disappear."),
            ("1", "Correct failing contrast tokens and verify all status/focus states.", "Automated contrast checks plus manual review meet applicable WCAG AA thresholds."),
            ("2", "Add reduced-motion behavior for spinner and smooth scrolling.", "prefers-reduced-motion users see no continuous spinner animation or animated scroll."),
            ("2", "Connect tooltip descriptions and rating slider values to their controls.", "Screen reader announces each explanation and current slider value."),
            ("2", "Add scopes/caption or aria-label to the past-ratings table.", "Table headers and relationships are announced correctly."),
            ("2", "Increase small custom tooltip hit targets and test responsive popovers.", "Controls have usable touch/keyboard targets and do not cover or trap content."),
            ("3", "Normalize decorative emoji treatment and add operation-completion messages.", "Screen-reader output is concise; success/error states are announced."),
            ("3", "Add automated accessibility checks to CI and document exceptions.", "axe/Lighthouse scan is run against representative authenticated and unauthenticated pages."),
        ],
        widths=[0.65, 3.45, 3.0],
        font_size=8,
    )

    heading("4.1 Suggested implementation sequence", level=2)
    for item in [
        "Establish shared accessibility utilities in base.html and style.css: skip link, focus target, reduced-motion rules, error-summary styles, and screen-reader status utility.",
        "Refactor the application wizard to use a single accessible step controller shared by caregiver and client forms.",
        "Refactor custom navigation/disclosure controls to use native buttons/details where possible instead of hover-dependent behavior.",
        "Update form rendering helpers so every field can receive described-by, invalid, and error-id attributes consistently.",
        "Run automated scans and keyboard checks, then conduct moderated user testing with disabled users before closing findings.",
    ]:
        numbered(item)

    page_break()

    # ---------------------------------------------------------------------
    # 5. VALIDATION AND MEASUREMENT PROCEDURE
    # ---------------------------------------------------------------------
    heading("5. Accessibility Validation and Measurement Procedure")
    body(
        "Accessibility should be measured as an ongoing product quality process rather than "
        "a one-time checklist. The following procedure is recommended for each material UI "
        "release and can be included in the project's iterative deploy, evaluate, refine, "
        "and retest process."
    )
    table(
        ["Measure", "Procedure", "Target / evidence"],
        [
            ("Automated rule violations", "Run axe-core or equivalent against public pages and representative authenticated pages in CI and before releases.", "Zero critical/serious violations; all accepted exceptions documented with owner and due date."),
            ("Keyboard task completion", "Complete login, application, schedule/rating, match review, and stability-flag workflows without a mouse.", "100% of critical workflows completable; no keyboard trap; visible focus at every step."),
            ("Screen-reader task completion", "Test with at least one mainstream desktop screen reader and one mobile screen reader where applicable.", "Users can identify page, headings, controls, errors, current wizard step, table relationships, and dynamic status."),
            ("Contrast and zoom", "Check text/UI states with a contrast analyzer and manually at 200%/400% zoom and 320px width.", "Applicable WCAG AA contrast thresholds met; no loss of essential content/functionality."),
            ("Reduced motion", "Test with prefers-reduced-motion enabled and disabled.", "No disorienting animation; loading and step changes remain understandable."),
            ("User experience", "Conduct moderated testing with disabled people who direct their own care, caregivers, PAS specialists, and ILC staff.", "Record task success, barriers, time on task, and qualitative feedback; prioritize issues by impact."),
            ("Regression tracking", "Store scan output, test date, browser/AT versions, and unresolved findings with each release.", "Trend serious findings toward zero and compare release-over-release performance."),
        ],
        widths=[1.45, 3.35, 2.3],
        font_size=8,
    )

    heading("5.1 Accessibility metrics for grant reporting", level=2)
    for item in [
        "Percentage of representative pages passing automated critical/serious checks.",
        "Percentage of critical user workflows completed by keyboard-only testers.",
        "Percentage of tested form fields with correctly announced labels, instructions, and errors.",
        "Number of unresolved Priority 1 accessibility findings per release.",
        "Percentage of disabled-user testing participants who can complete the assigned workflow without staff intervention.",
        "Median time to resolve a confirmed accessibility defect and percentage of defects re-tested successfully.",
    ]:
        bullet(item)

    heading("5.2 Privacy and respectful participation", level=2)
    body(
        "Accessibility testing will be conducted with informed consent and with respect for "
        "participants' communication preferences, access needs, and privacy. Test accounts and "
        "synthetic or de-identified records should be used whenever possible. Feedback should be "
        "recorded in the minimum detail needed to improve the product, and participants should be "
        "able to request an accommodation, alternate communication method, or human assistance "
        "without being penalized for doing so."
    )

    page_break()

    # ---------------------------------------------------------------------
    # 6. GRANT-READY STATEMENT
    # ---------------------------------------------------------------------
    heading("6. Grant-Ready Accessibility Statement")
    body(
        "CareWeb has completed a WCAG 2.1 AA-oriented source-code accessibility review of "
        "its shared layout, navigation, forms, dashboards, matching workflow, stability tools, "
        "and responsive styling. The review identified existing strengths, including semantic "
        "landmarks, explicit form labels, visible keyboard focus, accessible table structures, "
        "plain-language stability explanations, and status text that does not rely on color alone."
    )
    body(
        "The team also identified remediation priorities involving skip navigation, multi-step "
        "form announcements and validation, keyboard operation of dropdowns and disclosures, "
        "contrast in selected visual states, reduced-motion support, dynamic status announcements, "
        "and assistive-technology testing. These findings will be addressed through an iterative "
        "deploy, evaluate, refine, and retest process."
    )
    body(
        "Performance will be measured through automated accessibility scans, keyboard and screen-reader "
        "task completion, contrast and zoom checks, reduced-motion testing, and moderated usability "
        "sessions with disabled people who direct their own care, caregivers, PAS specialists, and "
        "Independent Living Center staff. Test results, unresolved findings, and remediation status "
        "will be tracked over time. Participation will use informed consent, de-identified test data "
        "where possible, and accommodations based on each participant's communication and access needs."
    )
    body(
        "This audit is a documented baseline and remediation plan; it should not be interpreted as a "
        "claim that the current prototype fully conforms to WCAG 2.1 AA until the recommended fixes "
        "and live assistive-technology testing are complete.",
        italic=True,
    )

    page_break()

    # ---------------------------------------------------------------------
    # APPENDIX A: FILE-BY-FILE NOTES
    # ---------------------------------------------------------------------
    heading("APPENDIX A — File-by-File Review Notes")
    table(
        ["File", "Observed accessibility controls", "Follow-up"],
        [
            ("templates/base.html", "lang, viewport, nav include, main/footer landmarks, AI status live region, burger state script", "Add skip link/main target, focus handling, reduced-motion/status completion."),
            ("templates/navbar.html", "Real burger button, aria-expanded/controls, nav label, POST logout button", "Refactor organization switcher; keyboard close/focus behavior."),
            ("templates/home.html", "Heading hierarchy, visible action labels, decorative badges aria-hidden", "Verify heading/navigation at zoom and screen-reader output."),
            ("templates/registration/login.html", "Explicit labels, CSRF, server-error rendering, visible submit button", "Add error summary/field descriptions and focus after failure."),
            ("templates/registry/caregiver_apply.html", "Explicit labels, section headings, visible errors, six-step workflow", "Make step state/validation/focus/announcements accessible."),
            ("templates/registry/schedule_entry_rate.html", "Explicit slider labels, visible current values, rating instructions", "Associate outputs/errors; add table scopes and accessible name."),
            ("templates/registry/org_dashboard.html", "Named table region, column scopes, status text, pagination include", "Test dense tables at zoom and with screen readers."),
            ("templates/matching/_match_table.html", "Named table/region, scope, text status labels, sr-only flag explanation", "Verify emoji output and popover/table reading order."),
            ("templates/matching/stability_detail.html", "Glossary, table scopes, labels, aria-expanded info buttons, status text", "Connect tooltip descriptions; verify keyboard/mobile popovers."),
            ("static/css/style.css", "Strong focus rules, readable base size, responsive rules, sr-only utility, status text/color pairing", "Fix contrast tokens, add reduced-motion, enlarge small targets."),
        ],
        widths=[1.75, 3.25, 2.1],
        font_size=8,
    )

    # ---------------------------------------------------------------------
    # APPENDIX B: REPRESENTATIVE CONTRAST RECORD
    # ---------------------------------------------------------------------
    heading("APPENDIX B — Representative Contrast Record")
    body(
        "The following ratios were calculated from the CSS hex tokens using the WCAG relative "
        "luminance formula. They are representative checks against white (#ffffff); actual "
        "rendered combinations and states must still be tested in the browser."
    )
    table(
        ["Foreground token", "Background", "Approx. ratio", "AA assessment"],
        [
            ("#7e6b8f (--app-primary)", "#ffffff", "4.79:1", "Passes normal text threshold"),
            ("#665575 (--app-primary-dark)", "#ffffff", "6.73:1", "Passes"),
            ("#3c6f99 (button/link blue)", "#ffffff", "5.35:1", "Passes normal text threshold"),
            ("#0f172a (heading/body dark)", "#ffffff", "17.85:1", "Passes"),
            ("#64748b (muted slate)", "#ffffff", "4.76:1", "Passes normal text threshold"),
            ("#94a3b8 (section label)", "#ffffff", "2.56:1", "Fails normal text threshold"),
            ("#da3e52 (--app-danger)", "#ffffff", "4.38:1", "Slightly below 4.5:1 for normal text"),
            ("#f2e94e (--app-warning)", "#ffffff", "1.27:1", "Do not use as light-text/background pair"),
            ("#1e293b (tooltip text)", "#f8fafc", "High contrast", "Passes representative check"),
        ],
        widths=[2.0, 1.1, 1.1, 2.5],
        font_size=8,
    )

    # Basic document metadata makes the generated file easier to identify.
    doc.core_properties.title = "CareWeb Accessibility Audit"
    doc.core_properties.subject = "WCAG 2.1 AA-oriented static accessibility review"
    doc.core_properties.author = "CareWeb Development Team"
    doc.core_properties.comments = "Generated from the CareWeb source-code accessibility review."
    doc.save(output_path)


if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output = os.path.join(base_dir, "CareWeb_Accessibility_Audit.docx")
    build_document(output)
    print(f"Document saved: {output}")